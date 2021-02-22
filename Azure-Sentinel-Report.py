import requests
import json
from requests.api import options
from docx import Document # pip3 install python-docx
from docx.shared import Inches
import matplotlib.pyplot as plt # sudo apt-get install python3-matplotlib
 

banner = """

Azure Sentinel Incident Report  POC

Usage: python3 Azure-Sentinel-Report.py 

"""

# Add the rquired fields

Azure_AD_Tenant = "Azure_AD_Tenant_HERE"
Client_ID = "Client_ID_HERE"
Client_Secret = "Client_Secret_HERE"
ResourceGroup = "ResourceGroup_HERE"
Workspace = "Workspace_HERE"
Subscription = "Subscription_ID"


# Get the Access Token

Url = "https://login.microsoftonline.com/"+Azure_AD_Tenant+"/oauth2/token"
headers = {'Content-Type': 'application/x-www-form-urlencoded'}
payload='grant_type=client_credentials&client_id='+ Client_ID+'&resource=https%3A%2F%2Fmanagement.azure.com&client_secret='+Client_Secret
response = requests.post(Url, headers=headers, data=payload).json()
Access_Token = response["access_token"]
print("[+] Access Token Received Successfully")


# Metric Queries
Url2= "https://management.azure.com/subscriptions/"+Subscription+"/resourceGroups/"+ResourceGroup+"/providers/Microsoft.OperationalInsights/workspaces/"+Workspace+"/api/query?api-version=2020-08-01"
payload2="{ \"query\": \"SecurityIncident | where TimeGenerated >= ago(30d) | summarize count()\"}"

Auth = 'Bearer '+Access_Token
headers2 = {
  'Authorization': Auth ,
  'Content-Type': 'text/plain'
}

response2 = requests.post(Url2, headers=headers2, data=payload2).json()
print("[+] Total Incidents were received Successfully")

payload3="{ \"query\": \"SecurityIncident | where TimeGenerated >= ago(30d) | summarize count() by Status\"}"
response3 = requests.post(Url2, headers=headers2, data=payload3).json()

payload4="{ \"query\": \"SecurityIncident | where TimeGenerated >= ago(30d) | summarize count() by Severity\"}"
response4 = requests.post(Url2, headers=headers2, data=payload4).json()

payload5="{ \"query\": \"Heartbeat | where TimeGenerated >= ago(30d) | summarize count() by OSType\"}"
response5 = requests.post(Url2, headers=headers2, data=payload5).json()


payload6="{ \"query\": \"Heartbeat | where TimeGenerated >= ago(30d) | summarize count() by RemoteIPCountry\"}"
response6 = requests.post(Url2, headers=headers2, data=payload6).json()


payload7="{ \"query\": \"SecurityIncident | where TimeGenerated >= ago(30d) | project Title, Description, Severity, FirstActivityTime, LastActivityTime\"}"
response7 = requests.post(Url2, headers=headers2, data=payload7).json()


payload8="{ \"query\": \"Heartbeat | where TimeGenerated >= ago(30d) | summarize count() by Computer\"}"
response8 = requests.post(Url2, headers=headers2, data=payload8).json()


payload9="{ \"query\": \"SecurityIncident | where TimeGenerated >= ago(30d) | summarize count() by bin(CreatedTime, 1h)\"}"
response9 = requests.post(Url2, headers=headers2, data=payload9).json()

# Metrics

Total_incidents = response2["Tables"][0]["Rows"][0][0]  
Total_assets = response8["Tables"][0]["Rows"]

#Incidents_by_Status
Inc1 = [] 
Stats = []
for i in range(len(response3["Tables"][0]["Rows"])):
    Inc1.append(response3["Tables"][0]["Rows"][i][1])
    Stats.append(response3["Tables"][0]["Rows"][i][0])


# Incidents_by_Severity
Inc2 = [] 
Sev = []
for i in range(len(response4["Tables"][0]["Rows"])):
    Inc2.append(response4["Tables"][0]["Rows"][i][1])
    Sev.append(response4["Tables"][0]["Rows"][i][0])


#Events_by_OS

Evs = [] 
OSs = []
for i in range(len(response5["Tables"][0]["Rows"])):
    Evs.append(response5["Tables"][0]["Rows"][i][1])
    OSs.append(response5["Tables"][0]["Rows"][i][0])


#Incidents_over_time
time = [] 
NB = []

for i in range(len(response9["Tables"][0]["Rows"])):
    time.append(response9["Tables"][0]["Rows"][i][0])
    NB.append(response9["Tables"][0]["Rows"][i][1])


#RemoteIP_Countries

Cnts = []
for i in range(len(response6["Tables"][0]["Rows"])):
    Cnts.append(response6["Tables"][0]["Rows"][i][0])

#Incident_details

Titles = []
Des = []
First = []
Last = []

for i in range(len(response7["Tables"][0]["Rows"])):
    Titles.append(response7["Tables"][0]["Rows"][i][0])
    Des.append(response7["Tables"][0]["Rows"][i][1])
    First.append(response7["Tables"][0]["Rows"][i][2])
    Last.append(response7["Tables"][0]["Rows"][i][3])

print("[+] The Incident details were Successfully loaded")

#OSsGRAPH

fig1 = plt.figure()
plt.pie(Evs,labels=OSs)
fig1.savefig("OSsGraph.png")

#Severity-GRAPH

fig2 = plt.figure() 
plt.pie(Inc2,labels=Sev)
fig2.savefig("Severity.png")

#Status Graph
fig3 = plt.figure() 
plt.pie(Inc1,labels=Stats)
fig3.savefig("Status.png")

#OverTime Graph 

fig4 = plt.figure() 
plt.bar(time,NB)
fig4.savefig("Overtime.png")


# Generate Word Report

document = Document()
document.add_heading('Azure Sentinel Report: Incidents Status ', 0)
document.add_heading('Agents Information', level=1)
document.add_paragraph('Number of Agents: '+ str(len(Total_assets)) )
#document.add_paragraph('Remote IP Countries: '+ len(Total_assets) )
document.add_paragraph('Agents by Operating System')
document.add_picture('OSsGraph.png',width=Inches(4))
document.add_heading('Incidents Status', level=1)
document.add_paragraph('Total Number of Incidents: '+ str(Total_incidents) )
document.add_paragraph('Incidents by Status')
document.add_picture('Status.png',width=Inches(4))
document.add_paragraph('Incidents by Severity')
document.add_picture('Severity.png',width=Inches(4))
document.add_paragraph('Incidents Over Time')
document.add_picture('Overtime.png',width=Inches(4))
document.add_heading('Incident Details', level=1)

table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Title'
hdr_cells[1].text = 'Description'
hdr_cells[2].text = 'First'
hdr_cells[3].text = 'Last'
for i in range(len(response7["Tables"][0]["Rows"])):
    row_cells = table.add_row().cells
    row_cells[0].text = Titles[i]
    row_cells[1].text = Des[i]
    row_cells[2].text = First[i]
    row_cells[3].text = Last[i]

document.save('Azure-Sentinel-Report.docx')
print("[+] The Report ( Azure-Sentinel-Report.docx ) was generated Successfully")


